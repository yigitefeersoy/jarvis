import threading, sys, os, platform, subprocess, webbrowser, datetime, asyncio, re, json
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from urllib.parse import urlparse, parse_qs
import numpy as np
import sounddevice as sd
import soundfile as sf
from scipy.signal import resample
import edge_tts
from groq import Groq 

GROQ_API_KEY = "API_KEYINIZ" #Buradan Bir API KEY Almanız Gerekmekte https://groq.com
MODEL = "openai/gpt-oss-120b"             
SES_ADI = "tr-TR-AhmetNeural"     
ORNEKLEME = 44100
SES_ESIGI = 0.04
MIKROFON = ""  #mikrofonunuzu otomatik seçmesi için adını yazabilirsiniz veya boş bırakabilirsiniz boş bırakıldığı zaman otomatik seçer.
PORT = 8777
ISLETIM = platform.system()


if GROQ_API_KEY and "BURAYA" not in GROQ_API_KEY:
    client = Groq(api_key=GROQ_API_KEY)
else:
    client 


def _masaustu():
    u = os.environ.get("USERPROFILE", os.path.expanduser("~"))
    for p in [os.path.join(u,"OneDrive","Masaüstü"), os.path.join(u,"OneDrive","Desktop"),
              os.path.join(u,"Desktop"), os.path.join(u,"Masaüstü")]:
        if os.path.isdir(p): return p
    return u

MASAUSTU = _masaustu()

def _ac_dosya(yol):
    if ISLETIM == "Windows": os.startfile(yol)
    else: subprocess.run(["open", yol])

def _temiz_ad(s):
    for c in '<>:"/\\|?*': s = s.replace(c, "")
    return (s.strip()[:50] or "belge")

durum = "baslatiliyor"
alt_yazi = "Hazirlaniyor..."
duydum = ""
calisiyor = True
whisper = None
MIK_INDEX = None

def giris_cihazlari():
    out, gorulen = [], set()
    for i, d in enumerate(sd.query_devices()):
        if d["max_input_channels"] > 0 and d["name"] not in gorulen:
            gorulen.add(d["name"]); out.append((i, d["name"]))
    return out

def mik_bul(parca):
    if not parca: return None
    for i, d in enumerate(sd.query_devices()):
        if d["max_input_channels"] > 0 and parca.lower() in d["name"].lower():
            return i
    return None

MIK_INDEX = mik_bul(MIKROFON)

def coz_ve_filtrele(saniye):
    ses = sd.rec(int(saniye * ORNEKLEME), samplerate=ORNEKLEME, channels=1, dtype="float32", device=MIK_INDEX)
    sd.wait()
    
    ses_seviyesi = float(np.abs(ses).max())
    
    if ses_seviyesi < SES_ESIGI:
        return None, ses_seviyesi
        
    hedef_ornekleme = 16000
    hedef_uzunluk = int(len(ses) * hedef_ornekleme / ORNEKLEME)
    ses_16k = resample(ses, hedef_uzunluk).astype(np.float32)
    
    sf.write("girdi.wav", ses_16k, hedef_ornekleme)
    segs, _ = whisper.transcribe("girdi.wav", language="tr")
    metin = " ".join(s.text for s in segs).strip().lower()
    
    gecersiz_kalıplar = ["görüşürüz", "bir sonraki videoda", "izlediğiniz için", "teşekkürler", "abone olun", "hoşça kalın"]
    if any(kalip in metin for kalip in gecersiz_kalıplar) and ses_seviyesi < (SES_ESIGI * 1.5):
        return "", ses_seviyesi
        
    return metin, ses_seviyesi

def konus(metin):
    try:
        asyncio.run(edge_tts.Communicate(metin, SES_ADI).save("cikti.mp3"))
        veri, hiz = sf.read("cikti.mp3"); sd.play(veri, hiz); sd.wait()
    except Exception as e:
        print("[konus HATA]", e)

UYGULAMA_HARITA = {
    "not defteri":"notepad","notepad":"notepad","defter":"notepad",
    "hesap makinesi":"calc","hesap makinasi":"calc","hesap":"calc",
    "paint":"mspaint","resim":"mspaint","boya":"mspaint",
    "ayarlar":"ms-settings:","dosya":"explorer","dosyalar":"explorer",
    "gezgin":"explorer","komut":"cmd","terminal":"cmd",
}
SITE_HARITA = {
    "youtube":"https://youtube.com","instagram":"https://instagram.com",
    "twitter":"https://twitter.com","x":"https://x.com","facebook":"https://facebook.com",
    "gmail":"https://mail.google.com","mail":"https://mail.google.com",
    "netflix":"https://netflix.com","spotify":"https://open.spotify.com",
    "whatsapp":"https://web.whatsapp.com","twitch":"https://twitch.tv",
    "github":"https://github.com","chatgpt":"https://chat.openai.com",
    "maps":"https://maps.google.com","harita":"https://maps.google.com",
    "reddit":"https://reddit.com","linkedin":"https://linkedin.com","disney":"https://disneyplus.com",
}
TARAYICI_KELIME = ["internet","tarayıcı","tarayici","browser","web sitesi"]

def uygulama_ac(ad):
    ad = ad.strip().lower().strip("'\". ")
    if ad in ("internet","tarayıcı","tarayici","browser","web","google") or any(k in ad for k in TARAYICI_KELIME):
        webbrowser.open("https://www.google.com"); return "Tarayiciyi aciyorum."
    if ad in SITE_HARITA:
        webbrowser.open(SITE_HARITA[ad]); return ad + " aciliyor."
    if ad in UYGULAMA_HARITA:
        hedef = UYGULAMA_HARITA[ad]
        try:
            if ISLETIM == "Windows":
                if hedef.endswith(":") or "://" in hedef: os.startfile(hedef)
                else: subprocess.Popen(hedef, shell=True)
            else: subprocess.run(["open","-a",hedef])
            return ad + " aciliyor."
        except Exception: return ad + " acilamadi."
    if " " not in ad and ad:
        webbrowser.open("https://www." + ad + ".com"); return ad + " sitesini aciyorum."
    webbrowser.open("https://www.google.com/search?q=" + ad); return ad + " icin arama actim."

def web_ara(sorgu):
    webbrowser.open("https://www.google.com/search?q=" + sorgu); return sorgu + " icin arama actim."

def youtube_ac(sorgu):
    webbrowser.open("https://www.youtube.com/results?search_query=" + sorgu); return sorgu + " icin YouTube actim."

def not_al(metin):
    with open("notlar.txt","a",encoding="utf-8") as f: f.write(metin + "\n")
    return "Notunu kaydettim."

def saat_soyle():
    return "Saat " + datetime.datetime.now().strftime("%H:%M") + "."

def ekrani_kilitle():
    if ISLETIM == "Windows": os.system("rundll32.exe user32.dll,LockWorkStation")
    else: os.system("pmset displaysleepnow")
    return "Ekrani kilitliyorum."

def klasor_olustur(isim):
    isim = _temiz_ad(isim) or "Yeni Klasor"
    yol = os.path.join(MASAUSTU, isim); os.makedirs(yol, exist_ok=True); _ac_dosya(yol)
    return isim + " klasorunu masaustunde olusturdum."

def belge_yaz(konu):
    konu = konu.strip() or "Belge"
    try:
        completion = client.chat.completions.create(
            model=MODEL,
            messages=[{"role": "user", "content": "Su konu hakkinda Turkce, duzenli, kisa 2 paragraf bir yazi yaz. Sadece yaziyi ver, baslik ekleme: " + konu}],
            temperature=0.5,
            max_completion_tokens=500,
            reasoning_effort="medium",
            stream=True
        )
        metin = ""
        for chunk in completion:
            metin += (chunk.choices[0].delta.content or "")
    except Exception: metin = ""
    
    try:
        from docx import Document
        d = Document(); d.add_heading(konu, 0)
        for par in (metin.split("\n") if metin else ["(icerik uretilemedi)"]):
            if par.strip(): d.add_paragraph(par.strip())
        dosya = os.path.join(MASAUSTU, _temiz_ad(konu) + ".docx")
        d.save(dosya); _ac_dosya(dosya)
        return konu + " hakkinda yaziyi yazip Word'de actim."
    except Exception as e:
        print("[belge_yaz HATA]", e); return "Belgeyi olusturamadim."

TALIMAT = (
    "Sen Jarvis adli bir asistansin. Istegi oku ve SADECE su komutlardan BIRINI tek satir ver. Aciklama ekleme:\n"
    "ARA|<sorgu>    -> internette aramak\n"
    "APP|<program>  -> program ACMAK (not defteri, hesap makinesi, tarayici, youtube...)\n"
    "YT|<sorgu>     -> youtube/muzik acmak\n"
    "NOT|<metin>    -> bilgi kaydetmek (sozunu AYNEN yaz)\n"
    "SAAT|          -> saati soylemek\n"
    "KILIT|         -> ekrani kilitlemek\n"
    "KLASOR|<isim>  -> masaustunde klasor olusturmak\n"
    "YAZI|<konu>    -> bir konuda yazi/word belgesi yazip acmak\n"
    "KONUS|<cevap>  -> normal sohbet (kisa, dogal, kuralina uygun Turkce cevap ver)\n"
    "ONEMLI: APP komutunda kullanicinin SOYLEDIGI hedefi aynen yaz. 'tarayici' sadece 'internet/tarayici ac' denirse.\n"
    "Ornekler:\n"
    "'not defterini ac' -> APP|notepad\n"
    "'youtube ac' -> APP|youtube\n"
    "'internet tarayicisini ac' -> APP|tarayici\n"
    "'masaustunde proje klasoru olustur' -> KLASOR|proje\n"
    "'kediler hakkinda yazi yaz' -> YAZI|kediler\n"
    "'kola almayi not al' -> NOT|kola al\n"
    "'hava nasil' -> ARA|bugun hava durumu\n"
    "Istek: "
)

def dusun_ve_yap(istek):
    try:
        completion = client.chat.completions.create(
            model=MODEL,
            messages=[{"role": "user", "content": TALIMAT + istek}],
            temperature=0,  # Komutların sekmemesi için kesinlik sıfır olmalı
            max_completion_tokens=150,
            top_p=1,
            reasoning_effort="medium",
            stream=True,
            stop=None
        )
        
        full_response = ""
        for chunk in completion:
            text_chunk = chunk.choices[0].delta.content or ""
            full_response += text_chunk
            print(text_chunk, end="")
        print() # Satır başı
        
        cevap = full_response.strip()
    except Exception as e:
        print("[Groq API HATA]", e)
        return "KONUS|Sisteme bağlanırken bir sorun oluştu."

    satir = cevap.splitlines()[0] if cevap else "KONUS|Anlayamadım efendim."
    komut, _, arg = satir.partition("|")
    komut = komut.strip().upper(); arg = arg.strip()
    print(f"[BEYIN] {satir}  -> KOMUT={komut} ARG={arg}")
    
    if komut.startswith("ARA"):                              return web_ara(arg)
    if komut.startswith("APP") or "ULAMA" in komut or "UYGU" in komut: return uygulama_ac(arg)
    if komut.startswith("YT") or "TUBE" in komut or "YOU" in komut:    return youtube_ac(arg)
    if komut.startswith("KLAS"):                             return klasor_olustur(arg)
    if komut.startswith("YAZ"):                              return belge_yaz(arg)
    if komut.startswith("NOT"):                              return not_al(arg)
    if komut.startswith("SAAT") or komut.startswith("SAT"):  return saat_soyle()
    if komut.startswith("KIL"):                              return ekrani_kilitle()
    return arg if arg else cevap

def beyin_dongusu():
    global durum, alt_yazi, duydum, whisper
    print("Jarvis Sistemi Resmi Groq SDK Aktif - Model:", MODEL)
    from faster_whisper import WhisperModel
    whisper = WhisperModel("small", device="cpu", compute_type="int8")
    
    while calisiyor:
        durum = "dinliyor"; alt_yazi = "Ortam dinleniyor..."
        istek, seviye = coz_ve_filtrele(4)
        
        if istek is None:
            continue
            
        if istek and len(istek.strip()) > 1:
            duydum = istek
            alt_yazi = istek; durum = "dusunuyor"
            try: cevap = dusun_ve_yap(istek)
            except Exception as e: cevap = "Sisteme ulasamadim."; print("[beyin]", e)
            durum = "konusuyor"; alt_yazi = cevap; konus(cevap)

SAYFA = """<!doctype html><html lang=tr><head><meta charset=utf-8>
<title>Jarvis HUD (GROQ SDK)</title>
<style>
html,body{margin:0;height:100%;background:radial-gradient(circle at 50% 45%,#05152d 0%,#01050f 70%);
overflow:hidden;font-family:'Segoe UI',Consolas,monospace;color:#cbe3ff;user-select:none}
canvas{position:fixed;inset:0}
#ust{position:fixed;top:14px;left:0;right:0;display:flex;justify-content:center;gap:8px;align-items:center;z-index:5}
#ust span{opacity:.5;font-size:12px;letter-spacing:.1em}
select{background:rgba(5,21,45,.75);color:#bfe0ff;border:1px solid #1f5078;border-radius:8px;
padding:6px 10px;font-size:13px;outline:none}
#baslik{position:fixed;top:8%;left:0;right:0;text-align:center;letter-spacing:.65em;font-weight:700;
font-size:26px;z-index:5;text-shadow:0 0 25px currentColor;color:#00a2ff}
#durumYazi{position:fixed;top:71%;left:0;right:0;text-align:center;font-weight:700;letter-spacing:.3em;
font-size:15px;z-index:5}
#altYazi{position:fixed;bottom:11%;left:0;right:0;text-align:center;font-size:20px;color:#e8f4ff;
z-index:5;padding:0 12%;text-shadow:0 0 12px rgba(0,162,255,.5)}
#duydum{position:fixed;bottom:5%;left:0;right:0;text-align:center;font-size:12px;color:#4f7390;z-index:5}
</style></head><body>
<canvas id=c></canvas>
<div id=ust><span>SİSTEM MİKROFONU</span><select id=mik></select></div>
<div id=baslik>J A R V I S</div>
<div id=durumYazi></div>
<div id=altYazi></div>
<div id=duydum></div>
<script>
const c=document.getElementById('c'),x=c.getContext('2d');
let durum='baslatiliyor',altYazi='',duydumY='';
const COL={baslatiliyor:[140,145,160],dinliyor:[0,162,255],dusunuyor:[0,255,234],konusuyor:[0,98,255]};
const ET={baslatiliyor:'SİSTEM YÜKLENİYOR',dinliyor:'AKILLI VAD KORUMASI AKTİF',dusunuyor:'GROQ AKIL YÜRÜTÜYOR',konusuyor:'Jarvis CEVAP VERİYOR'};
const SPD={baslatiliyor:1,dinliyor:2.4,dusunuyor:3.5,konusuyor:1.8};
function boyut(){c.width=innerWidth;c.height=innerHeight;}
addEventListener('resize',boyut);boyut();
function ciz(ts){
const w=c.width,h=c.height,cx=w/2,cy=h*0.45;
const col=COL[durum]||[0,162,255], c0=col.join(','), cs='rgb('+c0+')';
const spd=SPD[durum]||1, t=ts*0.001;
const puls=Math.sin(t*spd*2.2);
const R=Math.min(w,h)*0.14 + puls*8;
x.clearRect(0,0,w,h);
const g=x.createRadialGradient(cx,cy,R*0.3,cx,cy,R*3.4);
g.addColorStop(0,'rgba('+c0+',0.30)');g.addColorStop(0.45,'rgba('+c0+',0.08)');g.addColorStop(1,'rgba('+c0+',0)');
x.fillStyle=g;x.fillRect(0,0,w,h);
x.save();x.lineCap='round';x.shadowColor=cs;x.shadowBlur=28;x.strokeStyle=cs;
for(let ring=0;ring<3;ring++){
const rr=R*(1.45+ring*0.4), segs=8+ring*4, rot=t*0.4*(ring%2?-1:1)*(1+ring*0.3);
x.lineWidth=3.5-ring*0.8;
for(let s=0;s<segs;s++){const a0=rot+s*(2*Math.PI/segs),a1=a0+(2*Math.PI/segs)*0.55;
x.beginPath();x.arc(cx,cy,rr,a0,a1);x.stroke();}
}
x.shadowBlur=10;x.lineWidth=2;const tk=R*2.75;
for(let i=0;i<60;i++){const a=i*(2*Math.PI/60),uz=(i%5===0)?14:7;
x.globalAlpha=0.25+0.5*(i%5===0);
x.beginPath();x.moveTo(cx+Math.cos(a)*tk,cy+Math.sin(a)*tk);
x.lineTo(cx+Math.cos(a)*(tk+uz),cy+Math.sin(a)*(tk+uz));x.stroke();}
x.globalAlpha=1;x.restore();
const cg=x.createRadialGradient(cx,cy,2,cx,cy,R);
cg.addColorStop(0,'rgba(255,255,255,0.98)');cg.addColorStop(0.35,'rgba('+c0+',0.90)');
cg.addColorStop(1,'rgba('+c0+',0.10)');
x.fillStyle=cg;x.beginPath();x.arc(cx,cy,R,0,7);x.fill();
x.strokeStyle='rgba(255,255,255,0.5)';x.lineWidth=1.5;
x.beginPath();x.arc(cx,cy,R*0.74,0,7);x.stroke();
x.save();x.shadowColor=cs;x.shadowBlur=16;x.fillStyle=cs;
const N=50;
for(let i=0;i<N;i++){const a=t*0.7*spd+i*(2*Math.PI/N),rr=R*2.10+Math.sin(t*3+i)*8;
x.globalAlpha=0.3+0.6*Math.sin(t*4+i);
x.beginPath();x.arc(cx+Math.cos(a)*rr,cy+Math.sin(a)*rr,2.1,0,7);x.fill();}
x.restore();x.globalAlpha=1;
const dz=document.getElementById('durumYazi');dz.textContent=ET[durum]||'';dz.style.color=cs;dz.style.textShadow='0 0 15px '+cs;
document.getElementById('altYazi').textContent=altYazi;
document.getElementById('duydum').textContent=duydumY?('algılanan ses: '+duydumY):'';
requestAnimationFrame(ciz);
}
requestAnimationFrame(ciz);
async function durumCek(){try{const r=await fetch('/durum');const j=await r.json();
durum=j.durum;altYazi=j.yazi;duydumY=j.duydum;}catch(e){}}
setInterval(durumCek,200);durumCek();
async function mikYukle(){try{const r=await fetch('/mikrofonlar');const liste=await r.json();
const sel=document.getElementById('mik');sel.innerHTML='';
liste.forEach(m=>{const o=document.createElement('option');o.value=m.i;o.textContent=m.ad;
if(m.secili)o.selected=true;sel.appendChild(o);});
sel.onchange=()=>fetch('/mik?i='+sel.value,{method:'POST'});}catch(e){}}
mikYukle();
</script></body></html>"""

class Iste(BaseHTTPRequestHandler):
    def log_message(self, *a): pass
    def _gonder(self, kod, govde, tip="application/json"):
        self.send_response(kod); self.send_header("Content-Type", tip+"; charset=utf-8")
        self.end_headers(); self.wfile.write(govde.encode("utf-8"))
    def do_GET(self):
        yol = urlparse(self.path).path
        if yol == "/":
            self._gonder(200, SAYFA, "text/html")
        elif yol == "/durum":
            self._gonder(200, json.dumps({"durum":durum,"yazi":alt_yazi,"duydum":duydum}))
        elif yol == "/mikrofonlar":
            liste = [{"i":i,"ad":ad,"secili":(i==MIK_INDEX)} for i,ad in giris_cihazlari()]
            self._gonder(200, json.dumps(liste))
        else:
            self._gonder(404, "{}")
    def do_POST(self):
        global MIK_INDEX
        p = urlparse(self.path)
        if p.path == "/mik":
            q = parse_qs(p.query)
            try: MIK_INDEX = int(q.get("i",["0"])[0]); print("Mikrofon Değişti ->", MIK_INDEX)
            except Exception: pass
            self._gonder(200, "{}")
        else:
            self._gonder(404, "{}")

if __name__ == "__main__":
    threading.Thread(target=beyin_dongusu, daemon=True).start()
    sunucu = ThreadingHTTPServer(("127.0.0.1", PORT), Iste)
    webbrowser.open("http://127.0.0.1:%d" % PORT)
    print("Resmi Groq SDK Jarvis HUD Aktif: http://127.0.0.1:%d" % PORT)
    try: sunucu.serve_forever()
    except KeyboardInterrupt: calisiyor = False